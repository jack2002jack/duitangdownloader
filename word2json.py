import os
import docx
import json
files=[]
root="E:/data/新建文件夹/老师简介汇总"
for root,dirs,files_t in os.walk(root):
    files=files_t
print(files)
k={}
for file in files:
    doc=docx.Document(root+r"/"+file)
    for num in range(len(doc.paragraphs)):
        print(doc.paragraphs[num].text)
        name=file.split('.')[0]
        k[name]={"简介":"","研究方向":""}
fp=open("测试.json","w",encoding="utf-8")
json.dump(k,fp,ensure_ascii=False)